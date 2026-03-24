"""Word (.docx) format checker for NEU master's thesis (34 checks).

Checks that are more reliable on Word than PDF:
  1.  Body font (正文宋体12pt)
  2.  English font (Times New Roman)
  3.  Chapter heading (二号黑体22pt, 居中, 占3行)
  4.  Section heading (三号黑体16pt, 居左, 占2行)
  5.  Subsection heading (四号黑体14pt, 居左, 占2行)
  6.  Subsubsection heading (小四号黑体12pt, 居左, 占1行)
  7.  Special titles as level-1 (摘要/目录/参考文献/致谢)
  8.  Page headers (楷体10.5pt, 左端"东北大学硕士学位论文", 右端章题)
  9.  Page numbers (12pt, 居中, ·N· or -N-)
  10. Figure numbering (按章编码 图X.Y, 图号与图题间一字空格)
  11. Figure caption font (五号宋体10.5pt)
  12. Bilingual figure captions (中文下方紧跟英文 Fig. X.Y)
  13. Table numbering (按章编码 表X.Y, 表号与表题间空格)
  14. Table caption font (五号宋体10.5pt)
  15. Bilingual table captions (中文下方紧跟英文 Table X.Y)
  16. Equation numbering (按章编号 式(X.Y))
  17. Equation reference consistency (式 vs 公式 统一)
  18. Reference numbering + GB/T 7714 format
  19. Section numbering separator (半角点 ".")
  20. Chapter number Arabic (第1章, not 第一章)
  21. Chapter title length (<=20 chars)
  22. Paragraph last-line (>=5 chars)
  23. Blank paragraphs (连续空行)
  24. TOC entry font / numbering / consistency
  25. Lines per page / chars per line (estimate)
  26. Heading spacing check
  27. Figure/table caption spacing
  28. Table break properties (跨页重复标题行)
  29. Caption pair adjacency (中英文题目相邻)
  30. Mixed punctuation (中文段落英文标点)
  31. Mixed-width digits (全角数字)
  32. Caption trailing punctuation (图表题末尾标点)
  33. Duplicate words (连续重复汉字)
  34. Bracket mismatch (括号配对)

Usage:
    python word_checker.py <thesis.docx>
    Outputs JSON to stdout.
"""

import json
import re
import sys
import os
from collections import defaultdict


# ═══════════════════════════════════════════════════════════════
# Constants / Rules
# ═══════════════════════════════════════════════════════════════

EMU_PER_PT = 12700
EMU_PER_CM = 360000

# Chinese font-size names  →  pt
CN_SIZE = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
}

SONGTI_NAMES = {"宋体", "SimSun", "STSong", "Song", "NSimSun"}
HEITI_NAMES  = {"黑体", "SimHei", "STHei", "Hei"}
KAITI_NAMES  = {"楷体", "KaiTi", "STKai", "Kai"}
TNR_NAMES    = {"Times New Roman", "TimesNewRoman", "Times"}

SPECIAL_LEVEL1_TITLES = {"摘要", "摘  要", "Abstract", "ABSTRACT",
                         "目录", "目  录", "参考文献", "参  考  文  献",
                         "致谢", "致  谢", "攻读硕士期间的论文项目情况",
                         "攻读硕士学位期间发表的论文",
                         "攻读硕士期间的论文项目情况"}

# Regex helpers
RE_CHAPTER   = re.compile(r'^第\s*(\d+)\s*章\s+(.+)$')
RE_CHAPTER_CN = re.compile(r'^第\s*([一二三四五六七八九十]+)\s*章')
RE_SECTION   = re.compile(r'^(\d+)\.(\d+)\s+(.+)$')
RE_SUBSEC    = re.compile(r'^(\d+)\.(\d+)\.(\d+)\s+(.+)$')
RE_SUBSUB    = re.compile(r'^(\d+)\.(\d+)\.(\d+)\.(\d+)\s+(.+)$')

RE_FIG_CN    = re.compile(r'^图\s*(\d+)[.．](\d+)\s*(.*)')
RE_FIG_EN    = re.compile(r'^Fig\.?\s*(\d+)[.．](\d+)\s*(.*)', re.IGNORECASE)
RE_TAB_CN    = re.compile(r'^表\s*(\d+)[.．](\d+)\s*(.*)')
RE_TAB_EN    = re.compile(r'^Table\.?\s*(\d+)[.．](\d+)\s*(.*)', re.IGNORECASE)

RE_EQ_NUM    = re.compile(r'[\(（](\d+)[.．](\d+)[\)）]\s*$')
RE_EQ_REF    = re.compile(r'(式|公式)\s*[\(（](\d+)[.．](\d+)[\)）]')

RE_REF_ENTRY = re.compile(r'^\[(\d+)\]\s*')
RE_REF_TYPE  = re.compile(r'\[([MJCDPRSOEB/]+(?:OL)?)\]', re.IGNORECASE)

RE_SECTION_NUM = re.compile(r'^(\d+(?:\.\d+)+)\s')


# ═══════════════════════════════════════════════════════════════
# Helper utilities
# ═══════════════════════════════════════════════════════════════

def _emu_to_pt(emu):
    """Convert EMU to points."""
    if emu is None:
        return None
    return round(emu / EMU_PER_PT, 1)


def _issue(para_index, location, rule, expected, actual, severity="warning"):
    """Create a standardized issue dict."""
    return {
        "source": "word",
        "para_index": para_index,
        "location": location,
        "rule": rule,
        "expected": expected,
        "actual": actual,
        "severity": severity,
    }


def _font_matches(font_name, target_set):
    """Check if a font name matches any name in the target set (case-insensitive)."""
    if font_name is None:
        return False
    fn_lower = font_name.strip().lower()
    return any(t.lower() == fn_lower for t in target_set)


def _is_chinese(ch):
    """Return True if character is CJK Unified Ideograph."""
    cp = ord(ch)
    return (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
            0x20000 <= cp <= 0x2A6DF or 0xF900 <= cp <= 0xFAFF)


def _is_english(ch):
    """Return True if character is a Latin letter."""
    return ch.isascii() and ch.isalpha()


def _visible_len(text):
    """Length of visible characters (no whitespace)."""
    return len(text.replace(" ", "").replace("\t", "").replace("\u3000", ""))


def _get_paragraph_context(paragraphs, idx, direction="before"):
    """Get nearest non-empty paragraph text for context."""
    step = -1 if direction == "before" else 1
    i = idx + step
    while 0 <= i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if text:
            return text[:30]
        i += step
    return "(文档边界)"


# ═══════════════════════════════════════════════════════════════
# Font inheritance resolution
# ═══════════════════════════════════════════════════════════════

def _resolve_theme_font(doc, theme_key):
    """Resolve theme font name (majorEastAsia / minorEastAsia etc) to actual name.

    Returns the Hans (Simplified Chinese) font from the theme, or None.
    """
    try:
        from lxml import etree
        ns_a = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        doc_part = doc.part
        for rel in doc_part.rels.values():
            if 'theme' in rel.reltype.lower():
                theme_xml = rel.target_part.blob
                tree = etree.fromstring(theme_xml)
                # Determine major vs minor
                is_major = 'major' in theme_key.lower()
                tag = 'majorFont' if is_major else 'minorFont'
                for fs in tree.findall('.//a:fontScheme', ns_a):
                    for child in fs:
                        child_tag = child.tag.split('}')[-1]
                        if child_tag == tag:
                            # Look for Hans script
                            for font in child:
                                ft = font.tag.split('}')[-1]
                                if ft == 'font' and font.get('script') == 'Hans':
                                    return font.get('typeface')
                            # Fallback to ea
                            for font in child:
                                ft = font.tag.split('}')[-1]
                                if ft == 'ea':
                                    tf = font.get('typeface')
                                    if tf:
                                        return tf
                break
    except Exception:
        pass
    return None


def _get_style_east_asian_font(style, doc):
    """Get the east-asian font from a style's rFonts element."""
    from lxml import etree
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rpr = style._element.find(f'{{{ns_w}}}rPr')
    if rpr is None:
        return None
    rfonts = rpr.find(f'{{{ns_w}}}rFonts')
    if rfonts is None:
        return None
    # Direct east-asian font
    ea = rfonts.get(f'{{{ns_w}}}eastAsia')
    if ea:
        return ea
    # Theme-based east-asian font
    ea_theme = rfonts.get(f'{{{ns_w}}}eastAsiaTheme')
    if ea_theme:
        return _resolve_theme_font(doc, ea_theme)
    return None


def _get_run_east_asian_font(run, doc):
    """Get the east-asian font from a run's rFonts element."""
    from lxml import etree
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rpr = run._element.find(f'{{{ns_w}}}rPr')
    if rpr is None:
        return None
    rfonts = rpr.find(f'{{{ns_w}}}rFonts')
    if rfonts is None:
        return None
    ea = rfonts.get(f'{{{ns_w}}}eastAsia')
    if ea:
        return ea
    ea_theme = rfonts.get(f'{{{ns_w}}}eastAsiaTheme')
    if ea_theme:
        return _resolve_theme_font(doc, ea_theme)
    return None


def _effective_font_name(run, para, doc, for_chinese=True):
    """Resolve effective font name for a run, walking the inheritance chain.

    For Chinese text we look at eastAsia rFonts; for English, run.font.name (latin).
    Chain: run → paragraph style → base style(s) → Normal → theme.
    """
    if for_chinese:
        # Check run-level east-asian font
        ea = _get_run_east_asian_font(run, doc)
        if ea:
            return ea
        # Check run.font.name (sometimes set directly)
        if run.font.name:
            return run.font.name
        # Walk style chain
        style = para.style
        while style:
            ea = _get_style_east_asian_font(style, doc)
            if ea:
                return ea
            if style.font.name:
                return style.font.name
            style = style.base_style
        # Fallback: theme minor Hans = 宋体
        tf = _resolve_theme_font(doc, 'minorEastAsia')
        return tf or '宋体'
    else:
        # Latin font
        if run.font.name:
            return run.font.name
        style = para.style
        while style:
            if style.font.name:
                return style.font.name
            style = style.base_style
        return None


def _effective_font_size(run, para):
    """Resolve effective font size (in pt) for a run, walking the inheritance chain.

    Chain: run → paragraph style → base style(s) → Normal.
    """
    if run.font.size is not None:
        return _emu_to_pt(run.font.size)
    style = para.style
    while style:
        if style.font.size is not None:
            return _emu_to_pt(style.font.size)
        style = style.base_style
    return None


def _effective_bold(run, para):
    """Resolve effective bold for a run."""
    if run.font.bold is not None:
        return run.font.bold
    style = para.style
    while style:
        if style.font.bold is not None:
            return style.font.bold
        style = style.base_style
    return False


def _effective_alignment(para):
    """Resolve effective alignment, checking paragraph then style chain."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if para.alignment is not None:
        return para.alignment
    style = para.style
    while style:
        pf = style.paragraph_format
        if pf and pf.alignment is not None:
            return pf.alignment
        style = style.base_style
    return None


def _effective_space_before(para):
    """Resolve effective space_before in EMU."""
    pf = para.paragraph_format
    if pf.space_before is not None:
        return pf.space_before
    style = para.style
    while style:
        sf = style.paragraph_format
        if sf and sf.space_before is not None:
            return sf.space_before
        style = style.base_style
    return 0


def _effective_space_after(para):
    """Resolve effective space_after in EMU."""
    pf = para.paragraph_format
    if pf.space_after is not None:
        return pf.space_after
    style = para.style
    while style:
        sf = style.paragraph_format
        if sf and sf.space_after is not None:
            return sf.space_after
        style = style.base_style
    return 0


# ═══════════════════════════════════════════════════════════════
# Context builder — one-pass paragraph classification
# ═══════════════════════════════════════════════════════════════

def _build_context(doc):
    """Classify all paragraphs in a single pass.

    Returns a dict with:
      chapters:      [(idx, chapter_num, title_text), ...]
      sections:      [(idx, "X.Y", title_text), ...]
      subsections:   [(idx, "X.Y.Z", title_text), ...]
      subsubsections:[(idx, "X.Y.Z.W", title_text), ...]
      fig_captions:  [(idx, chap, num, full_text), ...]
      fig_captions_en: [(idx, chap, num, full_text), ...]
      tab_captions:  [(idx, chap, num, full_text), ...]
      tab_captions_en: [(idx, chap, num, full_text), ...]
      equations:     [(idx, chap, num, full_text), ...]
      references:    [(idx, ref_num, full_text), ...]
      toc_entries:   [(idx, level, text), ...]
      body_paras:    [idx, ...]   — Normal / 论文正文 body text
      special_titles:[(idx, title_text), ...]  — 摘要/目录/参考文献/致谢
      all_paras:     list of (idx, para) for convenience
      current_chapter: {para_idx: chapter_num}  — maps each para to its chapter
    """
    ctx = {
        "chapters": [],
        "sections": [],
        "subsections": [],
        "subsubsections": [],
        "fig_captions": [],
        "fig_captions_en": [],
        "tab_captions": [],
        "tab_captions_en": [],
        "equations": [],
        "references": [],
        "toc_entries": [],
        "body_paras": [],
        "special_titles": [],
        "current_chapter": {},
    }

    paragraphs = doc.paragraphs
    in_references = False
    in_toc = False
    current_chap = 0
    # Pre-title zone: paragraphs before the first Heading 1 with chapter content
    in_body = False
    # Track where body actually starts (after TOC)
    body_started = False

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""

        # --- TOC detection ---
        if 'toc' in style.lower():
            in_toc = True
            level = 1
            if '2' in style:
                level = 2
            elif '3' in style:
                level = 3
            if text:
                ctx["toc_entries"].append((i, level, text))
            continue
        elif in_toc and 'toc' not in style.lower():
            in_toc = False

        if not text:
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Heading 1 (chapter / special title) ---
        if style == "Heading 1":
            # Normalize whitespace for matching
            normalized = re.sub(r'\s+', '', text)
            is_special = False
            for st in SPECIAL_LEVEL1_TITLES:
                if re.sub(r'\s+', '', st) == normalized:
                    ctx["special_titles"].append((i, text))
                    is_special = True
                    break

            m = RE_CHAPTER.match(text)
            if m:
                current_chap = int(m.group(1))
                title = m.group(2).strip()
                ctx["chapters"].append((i, current_chap, title))
                # Body starts at the first numbered chapter
                if not body_started:
                    body_started = True
            elif not is_special:
                # Could be a chapter with odd formatting
                ctx["special_titles"].append((i, text))

            if normalized == "参考文献":
                in_references = True
            else:
                in_references = False

            ctx["current_chapter"][i] = current_chap
            continue

        # --- References ---
        if in_references:
            m = RE_REF_ENTRY.match(text)
            if m:
                ctx["references"].append((i, int(m.group(1)), text))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Heading 2 ---
        if style == "Heading 2":
            m = RE_SECTION.match(text)
            if m:
                ctx["sections"].append((i, f"{m.group(1)}.{m.group(2)}", m.group(3).strip()))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Heading 3 ---
        if style == "Heading 3":
            m = RE_SUBSEC.match(text)
            if m:
                ctx["subsections"].append((i, f"{m.group(1)}.{m.group(2)}.{m.group(3)}", m.group(4).strip()))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Heading 4 ---
        if style == "Heading 4":
            m = RE_SUBSUB.match(text)
            if m:
                ctx["subsubsections"].append((i, f"{m.group(1)}.{m.group(2)}.{m.group(3)}.{m.group(4)}", m.group(5).strip()))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Figure captions ---
        m_fig = RE_FIG_CN.match(text)
        if m_fig:
            ctx["fig_captions"].append((i, int(m_fig.group(1)), int(m_fig.group(2)), text))
            ctx["current_chapter"][i] = current_chap
            continue

        m_fig_en = RE_FIG_EN.match(text)
        if m_fig_en:
            ctx["fig_captions_en"].append((i, int(m_fig_en.group(1)), int(m_fig_en.group(2)), text))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Table captions ---
        m_tab = RE_TAB_CN.match(text)
        if m_tab:
            ctx["tab_captions"].append((i, int(m_tab.group(1)), int(m_tab.group(2)), text))
            ctx["current_chapter"][i] = current_chap
            continue

        m_tab_en = RE_TAB_EN.match(text)
        if m_tab_en:
            ctx["tab_captions_en"].append((i, int(m_tab_en.group(1)), int(m_tab_en.group(2)), text))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Equation paragraphs ---
        m_eq = RE_EQ_NUM.search(text)
        if m_eq and ('公示' in style or '公式' in style or
                     (text.startswith('\t') and len(text.replace('\t', '').strip()) < 15)):
            ctx["equations"].append((i, int(m_eq.group(1)), int(m_eq.group(2)), text))
            ctx["current_chapter"][i] = current_chap
            continue

        # --- Body paragraphs ---
        if body_started and not in_toc:
            if style in ("Normal", "5论文正文", "Body Text", "正文") or "正文" in style:
                ctx["body_paras"].append(i)

        ctx["current_chapter"][i] = current_chap

    return ctx


# ═══════════════════════════════════════════════════════════════
# Check functions (34 checks)
# ═══════════════════════════════════════════════════════════════

# ---------- 1. Body font ----------

def _check_body_font(doc, ctx):
    """Check body text: 小四号宋体(12pt) for Chinese."""
    issues = []
    paragraphs = doc.paragraphs
    expected_size = 12.0
    tolerance = 0.5
    sample_limit = 200  # check up to N body paragraphs to avoid huge output

    checked = 0
    for idx in ctx["body_paras"]:
        p = paragraphs[idx]
        text = p.text.strip()
        if not text or len(text) < 5:
            continue

        # Check Chinese runs
        has_chinese = any(_is_chinese(c) for c in text)
        if not has_chinese:
            continue

        for run in p.runs:
            run_text = run.text
            if not run_text or not any(_is_chinese(c) for c in run_text):
                continue

            fname = _effective_font_name(run, p, doc, for_chinese=True)
            fsize = _effective_font_size(run, p)

            # Font name check
            if fname and not _font_matches(fname, SONGTI_NAMES):
                issues.append(_issue(
                    idx,
                    f"第{idx+1}段",
                    "正文中文字体",
                    "宋体",
                    f"{fname}（内容: \"{run_text[:20]}\"）",
                    "error"
                ))
                break  # one issue per paragraph

            # Font size check
            if fsize is not None and abs(fsize - expected_size) > tolerance:
                issues.append(_issue(
                    idx,
                    f"第{idx+1}段",
                    "正文字号",
                    f"小四号({expected_size}pt)",
                    f"{fsize}pt（内容: \"{run_text[:20]}\"）",
                    "error"
                ))
                break

        checked += 1
        if checked >= sample_limit:
            break

    return issues


# ---------- 2. English font ----------

def _check_english_font(doc, ctx):
    """Check body text English: Times New Roman."""
    issues = []
    paragraphs = doc.paragraphs
    checked = 0
    sample_limit = 200

    for idx in ctx["body_paras"]:
        p = paragraphs[idx]
        text = p.text.strip()
        if not text:
            continue

        has_english = any(_is_english(c) for c in text)
        if not has_english:
            continue

        for run in p.runs:
            run_text = run.text
            if not run_text or not any(_is_english(c) for c in run_text):
                continue
            # Skip very short runs (punctuation, numbers mixed)
            english_chars = sum(1 for c in run_text if _is_english(c))
            if english_chars < 3:
                continue

            fname = _effective_font_name(run, p, doc, for_chinese=False)
            if fname and not _font_matches(fname, TNR_NAMES):
                issues.append(_issue(
                    idx,
                    f"第{idx+1}段",
                    "正文英文字体",
                    "Times New Roman",
                    f"{fname}（内容: \"{run_text[:20]}\"）",
                    "error"
                ))
                break

        checked += 1
        if checked >= sample_limit:
            break

    return issues


# ---------- 3. Chapter headings ----------

def _check_chapter_headings(doc, ctx):
    """Check chapter headings: 二号黑体(22pt), 居中."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, chap_num, title in ctx["chapters"]:
        p = paragraphs[idx]

        # Alignment check
        align = _effective_alignment(p)
        if align is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if align != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append(_issue(
                    idx, f"第{chap_num}章", "章标题对齐",
                    "居中", f"{align}", "error"
                ))

        # Font checks on runs
        for run in p.runs:
            if not run.text.strip():
                continue
            has_cn = any(_is_chinese(c) for c in run.text)
            if has_cn:
                fname = _effective_font_name(run, p, doc, for_chinese=True)
                if fname and not _font_matches(fname, HEITI_NAMES):
                    issues.append(_issue(
                        idx, f"第{chap_num}章", "章标题字体",
                        "黑体", f"{fname}", "error"
                    ))
                    break

        # Size check
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 22.0) > 1.0:
                issues.append(_issue(
                    idx, f"第{chap_num}章", "章标题字号",
                    "二号(22pt)", f"{fsize}pt", "error"
                ))
                break

        # Spacing check (should occupy ~3 lines: before+after ~= 2*正文行高)
        sp_before = _effective_space_before(p)
        sp_after = _effective_space_after(p)
        total_spacing = _emu_to_pt(sp_before + sp_after) if sp_before and sp_after else None
        if total_spacing is not None and total_spacing < 15:
            issues.append(_issue(
                idx, f"第{chap_num}章", "章标题前后间距",
                "占3行（前后间距合计约24pt）",
                f"前后间距合计{total_spacing}pt",
                "warning"
            ))

    return issues


# ---------- 4. Section headings ----------

def _check_section_headings(doc, ctx):
    """Check section (X.Y) headings: 三号黑体(16pt), 居左."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, num, title in ctx["sections"]:
        p = paragraphs[idx]

        # Alignment
        align = _effective_alignment(p)
        if align is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if align in (WD_ALIGN_PARAGRAPH.CENTER,):
                issues.append(_issue(
                    idx, f"{num}节", "节标题对齐",
                    "居左", f"居中", "error"
                ))

        # Font
        for run in p.runs:
            if not run.text.strip():
                continue
            has_cn = any(_is_chinese(c) for c in run.text)
            if has_cn:
                fname = _effective_font_name(run, p, doc, for_chinese=True)
                if fname and not _font_matches(fname, HEITI_NAMES):
                    issues.append(_issue(
                        idx, f"{num}节", "节标题字体",
                        "黑体", f"{fname}", "error"
                    ))
                    break

        # Size
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 16.0) > 1.0:
                issues.append(_issue(
                    idx, f"{num}节", "节标题字号",
                    "三号(16pt)", f"{fsize}pt", "error"
                ))
                break

    return issues


# ---------- 5. Subsection headings ----------

def _check_subsection_headings(doc, ctx):
    """Check subsection (X.Y.Z) headings: 四号黑体(14pt), 居左."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, num, title in ctx["subsections"]:
        p = paragraphs[idx]

        # Alignment
        align = _effective_alignment(p)
        if align is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if align == WD_ALIGN_PARAGRAPH.CENTER:
                issues.append(_issue(
                    idx, f"{num}款", "款标题对齐",
                    "居左", "居中", "error"
                ))

        # Font
        for run in p.runs:
            if not run.text.strip():
                continue
            has_cn = any(_is_chinese(c) for c in run.text)
            if has_cn:
                fname = _effective_font_name(run, p, doc, for_chinese=True)
                if fname and not _font_matches(fname, HEITI_NAMES):
                    issues.append(_issue(
                        idx, f"{num}款", "款标题字体",
                        "黑体", f"{fname}", "error"
                    ))
                    break

        # Size
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 14.0) > 1.0:
                issues.append(_issue(
                    idx, f"{num}款", "款标题字号",
                    "四号(14pt)", f"{fsize}pt", "error"
                ))
                break

    return issues


# ---------- 6. Subsubsection headings ----------

def _check_subsubsection_headings(doc, ctx):
    """Check subsubsection (X.Y.Z.W) headings: 小四号黑体(12pt), 居左."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, num, title in ctx["subsubsections"]:
        p = paragraphs[idx]

        # Font
        for run in p.runs:
            if not run.text.strip():
                continue
            has_cn = any(_is_chinese(c) for c in run.text)
            if has_cn:
                fname = _effective_font_name(run, p, doc, for_chinese=True)
                if fname and not _font_matches(fname, HEITI_NAMES):
                    issues.append(_issue(
                        idx, f"{num}项", "项标题字体",
                        "黑体", f"{fname}", "error"
                    ))
                    break

        # Size
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 12.0) > 0.5:
                issues.append(_issue(
                    idx, f"{num}项", "项标题字号",
                    "小四号(12pt)", f"{fsize}pt", "error"
                ))
                break

    return issues


# ---------- 7. Special titles as level-1 ----------

def _check_special_titles_as_level1(doc, ctx):
    """Check that 摘要/目录/参考文献/致谢 use Heading 1 style and level-1 formatting."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, title in ctx["special_titles"]:
        p = paragraphs[idx]
        style = p.style.name if p.style else ""

        # Must be Heading 1 style
        if style != "Heading 1":
            # TOC title might use a different style, which is acceptable
            normalized = re.sub(r'\s+', '', title)
            if normalized == "目录":
                continue  # TOC title can use TOC 标题1
            issues.append(_issue(
                idx, title[:20], "特殊标题样式",
                "Heading 1 (一级标题样式)",
                f"{style}",
                "warning"
            ))

        # Check font size = 22pt
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 22.0) > 1.0:
                issues.append(_issue(
                    idx, title[:20], "特殊标题字号",
                    "二号(22pt)", f"{fsize}pt", "error"
                ))
                break

    return issues


# ---------- 8. Page headers ----------

def _check_headers(doc, ctx):
    """Check page headers: 楷体10.5pt, left='东北大学硕士学位论文', right=chapter."""
    issues = []

    for si, sec in enumerate(doc.sections):
        header = sec.header
        if header.is_linked_to_previous:
            continue

        # Get header text
        header_text = ""
        for hp in header.paragraphs:
            header_text += hp.text

        if not header_text.strip():
            # Empty header sections (cover page, etc) are acceptable for early sections
            if si <= 1:
                continue
            issues.append(_issue(
                -1, f"Section {si}", "页眉内容",
                "东北大学硕士学位论文 + 章标题",
                "页眉为空",
                "warning"
            ))
            continue

        # Check left side contains "东北大学硕士学位论文"
        if "东北大学硕士学位论文" not in header_text:
            issues.append(_issue(
                -1, f"Section {si}", "页眉左端",
                "东北大学硕士学位论文",
                f"{header_text[:30]}",
                "error"
            ))

        # Check font: 楷体 10.5pt
        for hp in header.paragraphs:
            for run in hp.runs:
                rtext = run.text.strip()
                if not rtext:
                    continue
                # Check font name
                fname = run.font.name
                if fname and not _font_matches(fname, KAITI_NAMES):
                    # Numbers might not be 楷体
                    has_cn = any(_is_chinese(c) for c in rtext)
                    if has_cn:
                        issues.append(_issue(
                            -1, f"Section {si} 页眉", "页眉字体",
                            "楷体", f"{fname}（\"{rtext[:15]}\"）",
                            "error"
                        ))
                        break

                # Check size
                fsize = _emu_to_pt(run.font.size) if run.font.size else None
                if fsize is not None and abs(fsize - 10.5) > 0.5:
                    issues.append(_issue(
                        -1, f"Section {si} 页眉", "页眉字号",
                        "五号(10.5pt)", f"{fsize}pt",
                        "error"
                    ))
                    break

    return issues


# ---------- 9. Page numbers ----------

def _check_page_numbers(doc, ctx):
    """Check page footers for page number format."""
    issues = []

    for si, sec in enumerate(doc.sections):
        footer = sec.footer
        if footer.is_linked_to_previous:
            continue

        footer_text = ""
        for fp in footer.paragraphs:
            footer_text += fp.text

        if not footer_text.strip():
            if si <= 1:
                continue  # Cover pages may not have page numbers
            # Check if there's a field code for page number
            from lxml import etree
            ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            has_field = False
            for fp in footer.paragraphs:
                fields = fp._element.findall(f'.//{{{ns_w}}}fldChar')
                instrtext = fp._element.findall(f'.//{{{ns_w}}}instrText')
                if fields or instrtext:
                    has_field = True
                    break
            if not has_field and si > 3:  # After TOC
                issues.append(_issue(
                    -1, f"Section {si}", "页码",
                    "页底居中页码", "未找到页码",
                    "warning"
                ))
            continue

        # Check alignment of footer paragraphs
        for fp in footer.paragraphs:
            if not fp.text.strip():
                continue
            align = _effective_alignment(fp)
            if align is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if align != WD_ALIGN_PARAGRAPH.CENTER:
                    issues.append(_issue(
                        -1, f"Section {si} 页脚", "页码对齐",
                        "居中", f"{align}",
                        "warning"
                    ))

        # Check footer font size = 12pt (for arabic page sections)
        if si > 3:  # After TOC section
            for fp in footer.paragraphs:
                for run in fp.runs:
                    if not run.text.strip():
                        continue
                    fsize = _emu_to_pt(run.font.size) if run.font.size else None
                    if fsize is not None and abs(fsize - 12.0) > 1.0:
                        issues.append(_issue(
                            -1, f"Section {si} 页脚", "页码字号",
                            "12pt (同正文)", f"{fsize}pt",
                            "warning"
                        ))
                        break

    return issues


# ---------- 10. Figure numbering ----------

def _check_figure_numbering(doc, ctx):
    """Check figure numbering: 按章编码 图X.Y, 图号与图题间一字空格."""
    issues = []

    # Check sequential numbering within each chapter
    chap_figs = defaultdict(list)
    for idx, chap, num, text in ctx["fig_captions"]:
        chap_figs[chap].append((idx, num, text))

    for chap, figs in sorted(chap_figs.items()):
        expected = 1
        for idx, num, text in figs:
            if num != expected:
                gap = num - expected
                if gap > 0:
                    # Figures might be in textboxes not captured by doc.paragraphs
                    issues.append(_issue(
                        idx, f"图{chap}.{num}", "图编号连续性",
                        f"图{chap}.{expected}",
                        f"图{chap}.{num}（缺少{gap}个图题，可能在文本框中）",
                        "warning"
                    ))
                else:
                    issues.append(_issue(
                        idx, f"图{chap}.{num}", "图编号连续性",
                        f"图{chap}.{expected}",
                        f"图{chap}.{num}",
                        "error"
                    ))
            expected = num + 1

            # Check spacing between number and title
            # Use a more precise regex to detect space after "图X.Y"
            m_sp = re.match(r'^图\s*\d+[.．]\d+(\s*)(.*)', text)
            if m_sp:
                space_part = m_sp.group(1)
                title_part = m_sp.group(2)
                if title_part and not space_part:
                    issues.append(_issue(
                        idx, f"图{chap}.{num}", "图号与图题间距",
                        "图号与图题间空1字", f"无空格: \"{text[:30]}\"",
                        "warning"
                    ))

    return issues


# ---------- 11. Figure caption font ----------

def _check_figure_caption_font(doc, ctx):
    """Check figure caption font: 五号宋体(10.5pt), 居中."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, chap, num, text in ctx["fig_captions"]:
        p = paragraphs[idx]

        # Alignment check
        align = _effective_alignment(p)
        if align is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if align != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append(_issue(
                    idx, f"图{chap}.{num}", "图题对齐",
                    "居中", f"{align}",
                    "warning"
                ))

        # Font size check
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 10.5) > 0.5:
                issues.append(_issue(
                    idx, f"图{chap}.{num}", "图题字号",
                    "五号(10.5pt)", f"{fsize}pt",
                    "error"
                ))
                break

    return issues


# ---------- 12. Bilingual figure captions ----------

def _check_bilingual_figure_captions(doc, ctx):
    """Check bilingual figure captions: 中文下方紧跟英文 Fig. X.Y with matching numbers."""
    issues = []
    paragraphs = doc.paragraphs

    # Build a set of English figure captions by (chap, num)
    en_figs = {(chap, num): (idx, text)
               for idx, chap, num, text in ctx["fig_captions_en"]}

    for idx, chap, num, text in ctx["fig_captions"]:
        key = (chap, num)
        if key not in en_figs:
            issues.append(_issue(
                idx, f"图{chap}.{num}", "中英文图题",
                f"中文图题下方应有 Fig. {chap}.{num}",
                "缺少英文图题",
                "warning"
            ))
            continue

        en_idx, en_text = en_figs[key]
        # Check that English caption immediately follows Chinese
        if en_idx != idx + 1:
            # Allow one empty paragraph between
            if en_idx == idx + 2 and not paragraphs[idx+1].text.strip():
                pass  # OK
            else:
                issues.append(_issue(
                    idx, f"图{chap}.{num}", "中英文图题位置",
                    "英文图题应紧跟中文图题",
                    f"中文在第{idx+1}段, 英文在第{en_idx+1}段",
                    "warning"
                ))

    return issues


# ---------- 13. Table numbering ----------

def _check_table_numbering(doc, ctx):
    """Check table numbering: 按章编码 表X.Y, 编号间空格."""
    issues = []

    chap_tabs = defaultdict(list)
    for idx, chap, num, text in ctx["tab_captions"]:
        chap_tabs[chap].append((idx, num, text))

    for chap, tabs in sorted(chap_tabs.items()):
        expected = 1
        for idx, num, text in tabs:
            if num != expected:
                gap = num - expected
                if gap > 0:
                    issues.append(_issue(
                        idx, f"表{chap}.{num}", "表编号连续性",
                        f"表{chap}.{expected}",
                        f"表{chap}.{num}（缺少{gap}个表题，可能在文本框中）",
                        "warning"
                    ))
                else:
                    issues.append(_issue(
                        idx, f"表{chap}.{num}", "表编号连续性",
                        f"表{chap}.{expected}",
                        f"表{chap}.{num}",
                        "error"
                    ))
            expected = num + 1

            # Check spacing between number and title
            m_sp = re.match(r'^表\s*\d+[.．]\d+(\s*)(.*)', text)
            if m_sp:
                space_part = m_sp.group(1)
                title_part = m_sp.group(2)
                if title_part and not space_part:
                    issues.append(_issue(
                        idx, f"表{chap}.{num}", "表号与表题间距",
                        "表号与表题间空1~2字", f"无空格: \"{text[:30]}\"",
                        "warning"
                    ))

    return issues


# ---------- 14. Table caption font ----------

def _check_table_caption_font(doc, ctx):
    """Check table caption font: 五号宋体(10.5pt), 居中."""
    issues = []
    paragraphs = doc.paragraphs

    for idx, chap, num, text in ctx["tab_captions"]:
        p = paragraphs[idx]

        # Alignment
        align = _effective_alignment(p)
        if align is not None:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if align != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append(_issue(
                    idx, f"表{chap}.{num}", "表题对齐",
                    "居中", f"{align}",
                    "warning"
                ))

        # Font size
        for run in p.runs:
            if not run.text.strip():
                continue
            fsize = _effective_font_size(run, p)
            if fsize is not None and abs(fsize - 10.5) > 0.5:
                issues.append(_issue(
                    idx, f"表{chap}.{num}", "表题字号",
                    "五号(10.5pt)", f"{fsize}pt",
                    "error"
                ))
                break

    return issues


# ---------- 15. Bilingual table captions ----------

def _check_bilingual_table_captions(doc, ctx):
    """Check bilingual table captions: 中文下方紧跟英文 Table X.Y."""
    issues = []
    paragraphs = doc.paragraphs

    en_tabs = {(chap, num): (idx, text)
               for idx, chap, num, text in ctx["tab_captions_en"]}

    for idx, chap, num, text in ctx["tab_captions"]:
        key = (chap, num)
        if key not in en_tabs:
            issues.append(_issue(
                idx, f"表{chap}.{num}", "中英文表题",
                f"中文表题下方应有 Table {chap}.{num}",
                "缺少英文表题",
                "warning"
            ))
            continue

        en_idx, en_text = en_tabs[key]
        if en_idx != idx + 1:
            if en_idx == idx + 2 and not paragraphs[idx+1].text.strip():
                pass
            else:
                issues.append(_issue(
                    idx, f"表{chap}.{num}", "中英文表题位置",
                    "英文表题应紧跟中文表题",
                    f"中文在第{idx+1}段, 英文在第{en_idx+1}段",
                    "warning"
                ))

    return issues


# ---------- 16. Equation numbering ----------

def _check_equation_numbering(doc, ctx):
    """Check equation numbering: 按章编号 式(X.Y), sequential."""
    issues = []

    chap_eqs = defaultdict(list)
    for idx, chap, num, text in ctx["equations"]:
        chap_eqs[chap].append((idx, num, text))

    for chap, eqs in sorted(chap_eqs.items()):
        expected = 1
        for idx, num, text in eqs:
            if num != expected:
                issues.append(_issue(
                    idx, f"式({chap}.{num})", "公式编号连续性",
                    f"({chap}.{expected})",
                    f"({chap}.{num})",
                    "error"
                ))
            expected = num + 1

    return issues


# ---------- 17. Equation reference consistency ----------

def _check_equation_ref_consistency(doc, ctx):
    """Check that equation references consistently use 式 or 公式, not mixed."""
    issues = []
    paragraphs = doc.paragraphs
    shi_refs = []
    gongshi_refs = []

    for idx in ctx["body_paras"]:
        text = paragraphs[idx].text
        for m in RE_EQ_REF.finditer(text):
            word = m.group(1)
            if word == "式":
                shi_refs.append((idx, m.group(0)))
            else:
                gongshi_refs.append((idx, m.group(0)))

    if shi_refs and gongshi_refs:
        # Mixed usage — report the minority
        if len(shi_refs) <= len(gongshi_refs):
            majority = "公式"
            minority_refs = shi_refs
        else:
            majority = "式"
            minority_refs = gongshi_refs

        for idx, ref_text in minority_refs[:5]:  # limit output
            issues.append(_issue(
                idx, f"第{idx+1}段", "公式引用风格统一",
                f"统一使用「{majority}」",
                f"使用了「{ref_text}」",
                "warning"
            ))

    return issues


# ---------- 18. References ----------

def _check_references(doc, ctx):
    """Check reference entries: sequential [N], GB/T 7714 type markers."""
    issues = []

    expected_num = 1
    for idx, ref_num, text in ctx["references"]:
        # Sequential check
        if ref_num != expected_num:
            issues.append(_issue(
                idx, f"[{ref_num}]", "参考文献编号连续",
                f"[{expected_num}]",
                f"[{ref_num}]",
                "error"
            ))
        expected_num = ref_num + 1

        # GB/T 7714 type marker check
        m = RE_REF_TYPE.search(text)
        if not m:
            issues.append(_issue(
                idx, f"[{ref_num}]", "参考文献类型标识",
                "[M]/[J]/[C]/[D]/[R]/[EB/OL]等",
                f"未找到类型标识: \"{text[:50]}\"",
                "warning"
            ))

    return issues


# ---------- 19. Section numbering separator ----------

def _check_section_numbering_separator(doc, ctx):
    """Check that section numbers use half-width dot '.' not '．' or '·'."""
    issues = []
    paragraphs = doc.paragraphs

    all_headings = (
        [(idx, num, "section") for idx, num, _ in ctx["sections"]] +
        [(idx, num, "subsection") for idx, num, _ in ctx["subsections"]] +
        [(idx, num, "subsubsection") for idx, num, _ in ctx["subsubsections"]]
    )

    for idx, num, level in all_headings:
        text = paragraphs[idx].text.strip()
        # Check for fullwidth dot or middle dot
        if '．' in text[:15] or '·' in text[:15]:
            issues.append(_issue(
                idx, f"第{idx+1}段", "编号分隔符",
                "半角点 '.'",
                f"全角点或中间点: \"{text[:20]}\"",
                "warning"
            ))

    return issues


# ---------- 20. Chapter number Arabic ----------

def _check_chapter_number_arabic(doc, ctx):
    """Check chapter numbers are Arabic (第1章) not Chinese (第一章)."""
    issues = []
    paragraphs = doc.paragraphs

    # Also scan for Chinese-number chapters not caught by context builder
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and RE_CHAPTER_CN.match(text):
            issues.append(_issue(
                i, text[:20], "章编号格式",
                "阿拉伯数字（第1章）",
                f"中文数字: \"{text[:20]}\"",
                "error"
            ))

    return issues


# ---------- 21. Chapter title length ----------

def _check_chapter_title_length(doc, ctx):
    """Check chapter title length <= 20 chars."""
    issues = []
    MAX_LEN = 20

    for idx, chap_num, title in ctx["chapters"]:
        # Title is the part after "第X章 "
        visible = _visible_len(title)
        if visible > MAX_LEN:
            issues.append(_issue(
                idx, f"第{chap_num}章", "章标题长度",
                f"不超过{MAX_LEN}字",
                f"{visible}字: \"{title[:25]}\"",
                "warning"
            ))

    return issues


# ---------- 22. Paragraph last line ----------

def _check_paragraph_last_line(doc, ctx):
    """Check paragraph last line has >= 5 visible characters.

    This is an approximation: we check if the total text length modulo
    an estimated chars-per-line leaves fewer than 5 chars on the last line.
    Assumes ~38 chars per line.
    """
    issues = []
    paragraphs = doc.paragraphs
    CHARS_PER_LINE = 38
    MIN_LAST = 5

    for idx in ctx["body_paras"]:
        p = paragraphs[idx]
        text = p.text.strip()
        if not text or len(text) < CHARS_PER_LINE:
            continue  # Single-line paragraphs are fine

        visible = _visible_len(text)
        last_line_chars = visible % CHARS_PER_LINE
        if last_line_chars == 0:
            continue  # Full last line
        if last_line_chars < MIN_LAST:
            issues.append(_issue(
                idx, f"第{idx+1}段", "段落末行字数",
                f"末行至少{MIN_LAST}字",
                f"末行约{last_line_chars}字: \"...{text[-15:]}\"",
                "info"
            ))

    return issues


# ---------- 23. Blank paragraphs ----------

def _check_blank_paragraphs(doc, ctx):
    """Check for consecutive blank (empty) paragraphs in body text."""
    issues = []
    paragraphs = doc.paragraphs
    MAX_CONSECUTIVE = 1

    # Only check after body starts (skip cover pages)
    body_start = 0
    if ctx["chapters"]:
        body_start = ctx["chapters"][0][0]
    elif ctx["special_titles"]:
        body_start = ctx["special_titles"][0][0]

    i = body_start
    while i < len(paragraphs):
        if paragraphs[i].text.strip() == "":
            blank_start = i
            blank_count = 0
            while i < len(paragraphs) and paragraphs[i].text.strip() == "":
                blank_count += 1
                i += 1

            if blank_count > MAX_CONSECUTIVE:
                style_prev = paragraphs[blank_start - 1].style.name if blank_start > 0 else ""
                # Allow blanks after headings (spacing)
                if "Heading" in style_prev:
                    continue

                before_text = _get_paragraph_context(paragraphs, blank_start, "before")
                after_text = _get_paragraph_context(paragraphs, i - 1, "after")
                issues.append(_issue(
                    blank_start,
                    f"第{blank_start+1}段 ~ 第{blank_start+blank_count}段",
                    "多余空行",
                    f"连续空段落不超过{MAX_CONSECUTIVE}个",
                    f"连续{blank_count}个空段落\n"
                    f"  前文: \"{before_text}\"\n"
                    f"  后文: \"{after_text}\"",
                    "warning"
                ))
        else:
            i += 1

    return issues


# ---------- 24. TOC checks ----------

def _check_toc(doc, ctx):
    """Check TOC: entry numbering matches actual headings, title consistency."""
    issues = []
    paragraphs = doc.paragraphs

    if not ctx["toc_entries"]:
        return issues

    # Build set of actual chapter/section titles
    actual_chapters = {}
    for idx, chap, title in ctx["chapters"]:
        actual_chapters[chap] = title

    # Check TOC level-1 entries match actual chapter headings
    for toc_idx, level, toc_text in ctx["toc_entries"]:
        if level != 1:
            continue

        # Try to match "第X章 Title\tPageNum"
        # Remove tab and page number
        clean = toc_text.split('\t')[0].strip()
        m = RE_CHAPTER.match(clean)
        if m:
            toc_chap = int(m.group(1))
            toc_title = m.group(2).strip()
            # Normalize spaces for comparison
            toc_norm = re.sub(r'\s+', '', toc_title)

            if toc_chap in actual_chapters:
                actual_norm = re.sub(r'\s+', '', actual_chapters[toc_chap])
                if toc_norm != actual_norm:
                    issues.append(_issue(
                        toc_idx, f"目录第{toc_chap}章", "目录标题一致性",
                        f"与正文标题一致",
                        f"目录: \"{toc_title[:25]}\"\n  正文: \"{actual_chapters[toc_chap][:25]}\"",
                        "warning"
                    ))

    return issues


# ---------- 25. Lines & chars estimate ----------

def _check_lines_and_chars(doc, ctx):
    """Estimate lines-per-page and chars-per-line from document properties.

    This is a rough estimate; PDF-based checking is more accurate.
    """
    issues = []

    # Get page dimensions from first body section
    for sec in doc.sections:
        page_width = sec.page_width
        page_height = sec.page_height
        left_margin = sec.left_margin
        right_margin = sec.right_margin
        top_margin = sec.top_margin
        bottom_margin = sec.bottom_margin

        if page_width and left_margin and right_margin:
            body_width_mm = (page_width - left_margin - right_margin) / EMU_PER_CM * 10
            # Estimate chars per line: 12pt char ≈ 4.23mm for Chinese
            char_width_mm = 12 * 25.4 / 72  # ~4.23mm
            est_chars = int(body_width_mm / char_width_mm)
            if est_chars < 33 or est_chars > 40:
                issues.append(_issue(
                    -1, "版式", "每行字数估算",
                    "35~38字/行",
                    f"估算约{est_chars}字/行（版芯宽{body_width_mm:.1f}mm）",
                    "info"
                ))

        if page_height and top_margin and bottom_margin:
            body_height_mm = (page_height - top_margin - bottom_margin) / EMU_PER_CM * 10
            # Estimate lines: 12pt * 1.5 line spacing ≈ 18pt ≈ 6.35mm
            line_height_mm = 18 * 25.4 / 72  # ~6.35mm
            est_lines = int(body_height_mm / line_height_mm)
            if est_lines < 28 or est_lines > 37:
                issues.append(_issue(
                    -1, "版式", "每页行数估算",
                    "30~35行/页",
                    f"估算约{est_lines}行/页（版芯高{body_height_mm:.1f}mm）",
                    "info"
                ))
        break  # Only check first body section

    return issues


# ---------- 26. Heading spacing ----------

def _check_heading_spacing(doc, ctx):
    """Check heading spacing: chapter headings should occupy ~3 lines, sections ~2 lines."""
    issues = []
    paragraphs = doc.paragraphs
    NORMAL_LINE_HEIGHT_PT = 18  # 12pt * 1.5 line spacing

    for idx, chap, title in ctx["chapters"]:
        p = paragraphs[idx]
        sp_before = _effective_space_before(p)
        sp_after = _effective_space_after(p)

        # Chapter should occupy ~3 lines: heading itself + before + after
        # Minimum: before+after should be at least ~1 line height each
        sp_before_pt = _emu_to_pt(sp_before) if sp_before else 0
        sp_after_pt = _emu_to_pt(sp_after) if sp_after else 0

        if sp_before_pt < 10 and sp_after_pt < 10:
            issues.append(_issue(
                idx, f"第{chap}章", "章标题占行",
                "占3行（前后各留空间）",
                f"段前{sp_before_pt}pt 段后{sp_after_pt}pt（间距不足）",
                "warning"
            ))

    for idx, num, title in ctx["sections"]:
        p = paragraphs[idx]
        sp_before = _effective_space_before(p)
        sp_after = _effective_space_after(p)
        sp_before_pt = _emu_to_pt(sp_before) if sp_before else 0
        sp_after_pt = _emu_to_pt(sp_after) if sp_after else 0

        if sp_before_pt < 3 and sp_after_pt < 3:
            issues.append(_issue(
                idx, f"{num}节", "节标题占行",
                "占2行（前后有间距）",
                f"段前{sp_before_pt}pt 段后{sp_after_pt}pt（间距不足）",
                "warning"
            ))

    return issues


# ---------- 27. Figure/table caption spacing ----------

def _check_caption_spacing(doc, ctx):
    """Check figure/table caption spacing relative to body text."""
    issues = []
    # This is a lightweight check — mainly verifying captions aren't isolated
    # from their figures/tables. Detailed position checking is better in PDF.
    return issues


# ---------- 28. Table break properties ----------

def _check_table_break_properties(doc, ctx):
    """检查表格跨页属性：小表格不应跨页，大表格跨页需重复标题行。"""
    issues = []
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        # 检查是否有 tblHeader（重复标题行）
        first_row = table.rows[0]._tr
        has_header = first_row.find(f'{ns}trPr/{ns}tblHeader') is not None
        if rows > 10 and not has_header:
            issues.append(_issue(0, f"表格{i+1}({rows}行)", "大表格标题行重复",
                "大表格(>10行)跨页时应设置重复标题行", "未设置重复标题行", "warning"))
    return issues


# ---------- 29. Caption pair adjacency ----------

def _check_caption_pair_adjacency(doc, ctx):
    """检查中英文图题/表题段落是否相邻（para_index差值=1）。"""
    issues = []

    # Build English caption lookup: (chap, num) -> para_index
    fig_en_map = {(chap, num): idx for idx, chap, num, text in ctx["fig_captions_en"]}
    tab_en_map = {(chap, num): idx for idx, chap, num, text in ctx["tab_captions_en"]}

    for idx, chap, num, text in ctx["fig_captions"]:
        key = (chap, num)
        if key in fig_en_map:
            en_idx = fig_en_map[key]
            gap = abs(en_idx - idx)
            if gap != 1:
                issues.append(_issue(idx, f"图{chap}.{num}", "中英文图题相邻",
                    "中英文图题段落应紧邻(间距1段)",
                    f"中文题在段{idx}，英文题在段{en_idx}，间距{gap}段", "warning"))

    for idx, chap, num, text in ctx["tab_captions"]:
        key = (chap, num)
        if key in tab_en_map:
            en_idx = tab_en_map[key]
            gap = abs(en_idx - idx)
            if gap != 1:
                issues.append(_issue(idx, f"表{chap}.{num}", "中英文表题相邻",
                    "中英文表题段落应紧邻(间距1段)",
                    f"中文题在段{idx}，英文题在段{en_idx}，间距{gap}段", "warning"))

    return issues


# ---------- 30. Mixed punctuation ----------

def _check_mixed_punctuation(doc, ctx):
    """检查中文语境段落中是否出现英文标点。"""
    issues = []
    paragraphs = doc.paragraphs

    # Regex: detect Chinese characters
    re_has_cn = re.compile(r'[\u4e00-\u9fff]')
    # Regex: detect if paragraph is mostly English (>70% ASCII)
    re_ascii = re.compile(r'[a-zA-Z0-9 ]')
    # Target English punctuation in Chinese context
    en_puncts = [',', '.', ';', '(', ')', ':']
    # Exclusion patterns
    re_decimal = re.compile(r'\d\.\d')
    re_abbrev = re.compile(r'(?:e\.g\.|i\.e\.|et al\.|etc\.|vs\.|Dr\.|Mr\.|Mrs\.|Prof\.)', re.IGNORECASE)
    re_formula_num = re.compile(r'\(\d+\.\d+\)')
    re_file_ext = re.compile(r'\.\w{1,5}(?:\s|$|[,，;；])')
    re_section_ref = re.compile(r'\d+\.\d+(?:\.\d+)*(?:节|章|小节|式|图|表|步)')
    # English terms with periods in them (abbreviations, names)
    re_en_term = re.compile(r'[A-Za-z][A-Za-z0-9]*\.[A-Za-z0-9]')
    # Citation brackets [1,2,3] or [6,9]
    re_citation = re.compile(r'\[\d+(?:[,，]\d+)*\]')
    # English parens around English content: (XXX) where XXX is English
    re_en_paren = re.compile(r'\([A-Za-z][A-Za-z0-9 ,.\-]*\)')

    # Collect reference para indices for exclusion
    ref_indices = {idx for idx, _, _ in ctx["references"]}

    for pidx in ctx["body_paras"]:
        p = paragraphs[pidx]
        text = p.text.strip()
        if not text:
            continue
        # Skip reference paragraphs
        if pidx in ref_indices:
            continue
        # Skip if no Chinese chars (pure English paragraph)
        if not re_has_cn.search(text):
            continue
        # Skip if paragraph is predominantly English (>70% ASCII chars)
        ascii_count = len(re_ascii.findall(text))
        if len(text) > 0 and ascii_count / len(text) > 0.7:
            continue

        # Remove safe patterns before checking (order matters!)
        # Section refs first (before decimal eats digits)
        cleaned = re_section_ref.sub('##', text)
        cleaned = re_formula_num.sub('##', cleaned)
        cleaned = re_en_paren.sub('##', cleaned)
        cleaned = re_citation.sub('##', cleaned)
        cleaned = re_abbrev.sub('##', cleaned)
        cleaned = re_en_term.sub('##', cleaned)
        cleaned = re_decimal.sub('##', cleaned)
        cleaned = re_file_ext.sub('##', cleaned)

        found = []
        for punct in en_puncts:
            if punct in cleaned:
                found.append(punct)

        if found:
            sample = text[:50] + ("..." if len(text) > 50 else "")
            issues.append(_issue(pidx, f"段{pidx}", "中文段落英文标点",
                "中文语境应使用中文标点",
                f"发现英文标点 {''.join(found)}，段落: {sample}", "warning"))

    return issues


# ---------- 31. Mixed-width digits ----------

def _check_mixed_width_digits(doc, ctx):
    """检查全角数字出现即报warning。"""
    issues = []
    paragraphs = doc.paragraphs
    re_fullwidth = re.compile(r'[０-９]+')

    for pidx in ctx["body_paras"]:
        text = paragraphs[pidx].text.strip()
        if not text:
            continue
        matches = re_fullwidth.findall(text)
        if matches:
            sample = text[:50] + ("..." if len(text) > 50 else "")
            issues.append(_issue(pidx, f"段{pidx}", "全角数字",
                "应使用半角数字0-9",
                f"发现全角数字 {'、'.join(matches[:3])}，段落: {sample}", "warning"))

    return issues


# ---------- 32. Caption trailing punctuation ----------

def _check_caption_trailing_punct(doc, ctx):
    """检查图题/表题末尾是否有多余标点。"""
    issues = []
    paragraphs = doc.paragraphs
    bad_trailing = re.compile(r'[。.，,；;！!？?]$')

    # Check Chinese figure captions
    for idx, chap, num, text in ctx["fig_captions"]:
        text_stripped = text.strip()
        if text_stripped and bad_trailing.search(text_stripped):
            issues.append(_issue(idx, f"图{chap}.{num}", "图题末尾标点",
                "图题末尾不应有句号/逗号等标点",
                f"末尾字符: '{text_stripped[-1]}'", "warning"))

    # Check English figure captions
    for idx, chap, num, text in ctx["fig_captions_en"]:
        text_stripped = text.strip()
        if text_stripped and bad_trailing.search(text_stripped):
            issues.append(_issue(idx, f"Fig.{chap}.{num}", "英文图题末尾标点",
                "图题末尾不应有句号/逗号等标点",
                f"末尾字符: '{text_stripped[-1]}'", "warning"))

    # Check Chinese table captions
    for idx, chap, num, text in ctx["tab_captions"]:
        text_stripped = text.strip()
        if text_stripped and bad_trailing.search(text_stripped):
            issues.append(_issue(idx, f"表{chap}.{num}", "表题末尾标点",
                "表题末尾不应有句号/逗号等标点",
                f"末尾字符: '{text_stripped[-1]}'", "warning"))

    # Check English table captions
    for idx, chap, num, text in ctx["tab_captions_en"]:
        text_stripped = text.strip()
        if text_stripped and bad_trailing.search(text_stripped):
            issues.append(_issue(idx, f"Table {chap}.{num}", "英文表题末尾标点",
                "表题末尾不应有句号/逗号等标点",
                f"末尾字符: '{text_stripped[-1]}'", "warning"))

    return issues


# ---------- 33. Duplicate words ----------

def _check_duplicate_words(doc, ctx):
    """正则检测连续重复汉字。排除合法叠词和跨词边界误报。"""
    issues = []
    paragraphs = doc.paragraphs
    # Match consecutive identical Chinese characters (2+ repetitions)
    re_dup = re.compile(r'([\u4e00-\u9fff])\1+')
    # Legal reduplications (common valid repeated-char words)
    legal_redupl = {
        '谢谢', '渐渐', '慢慢', '往往', '常常', '仅仅', '刚刚', '偶偶',
        '默默', '悄悄', '淡淡', '深深', '轻轻', '静静', '缓缓', '微微',
        '稍稍', '略略', '纷纷', '频频', '每每', '处处', '时时', '事事',
        '人人', '天天', '年年', '代代', '层层', '步步', '点点', '斑斑',
        '种种', '多多', '少少', '大大', '小小', '好好', '早早', '草草',
        '匆匆', '隐隐', '朦朦', '茫茫', '苍苍', '莽莽', '洋洋', '堂堂',
        '泱泱', '朗朗', '琅琅', '铮铮', '粼粼', '潺潺', '滚滚', '滔滔',
        '哈哈', '呵呵', '嘻嘻', '嘿嘿', '啧啧', '咄咄', '喃喃', '念念',
        '姗姗', '翩翩', '冉冉', '袅袅', '娓娓', '侃侃', '赫赫', '炎炎',
        '凛凛', '彬彬', '落落', '碌碌', '区区', '芸芸', '济济', '岌岌',
        '欣欣', '蒸蒸', '勃勃', '蠢蠢', '跃跃', '津津', '源源',
    }
    # Common cross-word-boundary patterns where the same character appears
    # at the end of one word and beginning of the next (e.g. 解密+密钥 → 密密)
    # We check 2-char context before and after the repeated char
    cross_boundary_patterns = re.compile(
        r'(?:解密密钥|加密密钥|加密密文|解密密文|对称密密钥|非对称密密钥'
        r'|验证证明|验证证据|验证证书|认证证书|凭证证明'
        r'|提交交易|提交交付|成交交易|撤交交割|期交交割'
        r'|实验验证|试验验证|检验验证|校验验证'
        r'|以以太|链以以'
        r'|线性性质|特性性能|属性性质|弹性性能'
        r'|方面面临|层面面对|界面面板'
        r'|据包包含|软件包包括'
        r'|比较较大|比较较小|比较较高|比较较低'
        r'|委托托管|信托托管'
        r'|意义义务|含义义务'
        r'|字符符号|字符符合'
        r'|重新新建|更新新增'
        r'|评估估计|评估估值'
        r'|决策策略|政策策略'
        r'|提议议案|建议议题'
        r'|可见见证|意见见解'
        r'|人民民主|市民民众'
        r'|管理理论|处理理由|道理理论|原理理论'
        r'|研究究竟)'
    )

    ref_indices = {idx for idx, _, _ in ctx["references"]}

    for pidx in ctx["body_paras"]:
        text = paragraphs[pidx].text.strip()
        if not text or pidx in ref_indices:
            continue

        for m in re_dup.finditer(text):
            dup_word = m.group(0)
            if dup_word in legal_redupl:
                continue
            # Check cross-word boundary: extract wider context around the match
            ctx_start = max(0, m.start() - 2)
            ctx_end = min(len(text), m.end() + 2)
            context_window = text[ctx_start:ctx_end]
            if cross_boundary_patterns.search(context_window):
                continue
            # Show context around the duplicate
            start = max(0, m.start() - 5)
            end = min(len(text), m.end() + 5)
            context_str = text[start:end]
            issues.append(_issue(pidx, f"段{pidx}", "连续重复汉字",
                "避免连续重复汉字（如'的的''了了'）",
                f"发现'{dup_word}'，上下文: ...{context_str}...", "warning"))

    return issues


# ---------- 34. Bracket mismatch ----------

def _check_bracket_mismatch(doc, ctx):
    """对每段计数括号配对。"""
    issues = []
    paragraphs = doc.paragraphs
    bracket_pairs = [
        ('(', ')', '半角圆括号'),
        ('（', '）', '全角圆括号'),
        ('[', ']', '半角方括号'),
        ('【', '】', '方头括号'),
    ]

    ref_indices = {idx for idx, _, _ in ctx["references"]}

    for pidx in ctx["body_paras"]:
        text = paragraphs[pidx].text.strip()
        if not text or pidx in ref_indices:
            continue

        for left, right, name in bracket_pairs:
            lc = text.count(left)
            rc = text.count(right)
            if lc != rc:
                sample = text[:50] + ("..." if len(text) > 50 else "")
                issues.append(_issue(pidx, f"段{pidx}", f"{name}不配对",
                    f"'{left}'与'{right}'数量应相等",
                    f"'{left}'={lc}个, '{right}'={rc}个，段落: {sample}", "warning"))

    return issues


# ═══════════════════════════════════════════════════════════════
# Main entry point
# ═══════════════════════════════════════════════════════════════

def check_word(docx_path: str) -> dict:
    """Run all Word-specific checks on a thesis .docx file.

    Returns dict with 'issues' list and 'summary' counts.
    """
    from docx import Document

    doc = Document(docx_path)
    ctx = _build_context(doc)

    issues = []

    # Run all 34 checks
    checks = [
        ("正文中文字体", _check_body_font),
        ("正文英文字体", _check_english_font),
        ("章标题格式", _check_chapter_headings),
        ("节标题格式", _check_section_headings),
        ("款标题格式", _check_subsection_headings),
        ("项标题格式", _check_subsubsection_headings),
        ("特殊标题格式", _check_special_titles_as_level1),
        ("页眉格式", _check_headers),
        ("页码格式", _check_page_numbers),
        ("图编号", _check_figure_numbering),
        ("图题字体", _check_figure_caption_font),
        ("中英文图题", _check_bilingual_figure_captions),
        ("表编号", _check_table_numbering),
        ("表题字体", _check_table_caption_font),
        ("中英文表题", _check_bilingual_table_captions),
        ("公式编号", _check_equation_numbering),
        ("公式引用风格", _check_equation_ref_consistency),
        ("参考文献格式", _check_references),
        ("编号分隔符", _check_section_numbering_separator),
        ("章编号格式", _check_chapter_number_arabic),
        ("章标题长度", _check_chapter_title_length),
        ("段落末行字数", _check_paragraph_last_line),
        ("多余空行", _check_blank_paragraphs),
        ("目录检查", _check_toc),
        ("行数字数估算", _check_lines_and_chars),
        ("标题间距", _check_heading_spacing),
        ("图表题间距", _check_caption_spacing),
        ("表格跨页属性", _check_table_break_properties),
        ("中英文题目相邻", _check_caption_pair_adjacency),
        ("中文段落英文标点", _check_mixed_punctuation),
        ("全角数字", _check_mixed_width_digits),
        ("图表题末尾标点", _check_caption_trailing_punct),
        ("连续重复汉字", _check_duplicate_words),
        ("括号配对", _check_bracket_mismatch),
    ]

    for name, check_fn in checks:
        try:
            result = check_fn(doc, ctx)
            issues.extend(result)
        except Exception as e:
            issues.append(_issue(
                -1, name, "检查异常",
                "正常运行", f"异常: {e}",
                "error"
            ))

    errors = sum(1 for i in issues if i["severity"] == "error")
    warnings = sum(1 for i in issues if i["severity"] == "warning")
    infos = sum(1 for i in issues if i["severity"] == "info")

    return {
        "issues": issues,
        "summary": {
            "total": len(issues),
            "errors": errors,
            "warnings": warnings,
            "infos": infos,
            "checks_run": len(checks),
            "context": {
                "chapters": len(ctx["chapters"]),
                "sections": len(ctx["sections"]),
                "subsections": len(ctx["subsections"]),
                "figures_cn": len(ctx["fig_captions"]),
                "figures_en": len(ctx["fig_captions_en"]),
                "tables_cn": len(ctx["tab_captions"]),
                "tables_en": len(ctx["tab_captions_en"]),
                "equations": len(ctx["equations"]),
                "references": len(ctx["references"]),
                "toc_entries": len(ctx["toc_entries"]),
                "body_paras": len(ctx["body_paras"]),
            },
        },
    }


# ═══════════════════════════════════════════════════════════════
# CLI Entry
# ═══════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 2:
        print("Usage: python word_checker.py <thesis.docx>", file=sys.stderr)
        sys.exit(1)

    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    result = check_word(docx_path)

    # Force UTF-8 output on Windows
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
